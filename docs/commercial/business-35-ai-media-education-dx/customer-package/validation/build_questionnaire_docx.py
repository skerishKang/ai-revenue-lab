#!/usr/bin/env python3
"""Generate the Business 35 diagnostic questionnaire DOCX (3 pages).

Commercial truth is consumed from the exact accepted Lane A revision via
``accepted_source`` (03-diagnostic-questionnaire.md). Q1-Q5 implement the
five V3.1 input dimensions as actual fillable fields (조직 유형 / 결과물
유형 / 병목 지점 / 현재 팀 규모 / AI 사용 상태); Q6-Q17 are flow, baseline,
governance and readiness detail.

- Title in one/two natural lines
- Page headers repeated per page
- Narrative questions get 2-3 answer lines
- Yes/No questions get real checkboxes; choice questions get selection cells
- No internal English status markers in the customer document
- Notice block at the bottom of the last page
- 3 pages maximum
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_BREAK

from pathlib import Path
import datetime
import sys
sys.path.insert(0, str(Path(__file__).resolve().parent))
from accepted_source import require_accepted_source, diagnostic_q1_q5  # noqa: E402
FIXED_DT = datetime.datetime(2026, 9, 3, 0, 0, 0)
OUT = Path(__file__).resolve().parent.parent / "Business35_Diagnostic_Questionnaire.docx"

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GRAY = RGBColor(0x55, 0x5A, 0x60)


def heading(doc, text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = NAVY
        run.font.name = "Malgun Gothic"
        run.font.size = Pt(13)
    return h


def narrative(doc, question, guide):
    p = doc.add_paragraph()
    r = p.add_run(question)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.name = "Malgun Gothic"
    p2 = doc.add_paragraph()
    r2 = p2.add_run(f"({guide})")
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = GRAY
    r2.font.name = "Malgun Gothic"
    for _ in range(2):
        line = doc.add_paragraph()
        lr = line.add_run("_" * 60)
        lr.font.size = Pt(11)
        lr.font.name = "Malgun Gothic"
        line.paragraph_format.space_after = Pt(8)


def short_answer(doc, question, guide):
    p = doc.add_paragraph()
    r = p.add_run(question)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.name = "Malgun Gothic"
    p2 = doc.add_paragraph()
    r2 = p2.add_run(f"({guide})")
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = GRAY
    r2.font.name = "Malgun Gothic"
    line = doc.add_paragraph()
    lr = line.add_run("_" * 60)
    lr.font.size = Pt(11)
    lr.font.name = "Malgun Gothic"
    line.paragraph_format.space_after = Pt(8)


def yesno(doc, question, options):
    p = doc.add_paragraph()
    r = p.add_run(question)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.name = "Malgun Gothic"
    op = doc.add_paragraph()
    for label in options:
        run = op.add_run(f"  ☐  {label}")
        run.font.size = Pt(10.5)
        run.font.name = "Malgun Gothic"
    op.paragraph_format.space_after = Pt(8)


def choice(doc, question, options):
    p = doc.add_paragraph()
    r = p.add_run(question)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.name = "Malgun Gothic"
    op = doc.add_paragraph()
    for label in options:
        run = op.add_run(f"  ☐  {label}    ")
        run.font.size = Pt(10.5)
        run.font.name = "Malgun Gothic"
    op.paragraph_format.space_after = Pt(8)


def page_header(doc, title, page_no):
    ph = doc.add_paragraph()
    pr = ph.add_run("파디엠 · AI 업무전환 진단 질문지   ")
    pr.font.size = Pt(10)
    pr.font.bold = True
    pr.font.color.rgb = NAVY
    pr.font.name = "Malgun Gothic"
    pt = ph.add_run(title + "   ")
    pt.font.size = Pt(12)
    pt.font.bold = True
    pt.font.color.rgb = NAVY
    pt.font.name = "Malgun Gothic"
    prr = ph.add_run(f"[ {page_no} / 3 ]")
    prr.font.size = Pt(9)
    prr.font.color.rgb = GRAY
    prr.font.name = "Malgun Gothic"
    ph.paragraph_format.space_after = Pt(8)


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build():
    snapshot = require_accepted_source()
    q1_q5 = diagnostic_q1_q5(snapshot)  # fail-closed: must be the accepted five inputs
    assert [t for _, t in q1_q5] == ["조직 유형", "결과물 유형", "병목 지점", "현재 팀 규모", "AI 사용 상태"]
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(10.5)

    # Narrower margins to fit 3 pages
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # ---- Page 1 ----
    title = doc.add_heading("Business 35 · 파디엠 AI 미디어 업무전환 스튜디오", level=0)
    for run in title.runs:
        run.font.color.rgb = NAVY
        run.font.name = "Malgun Gothic"
        run.font.size = Pt(18)

    sub = doc.add_paragraph()
    sr = sub.add_run("고객 진단 질문지 · V3.1 다섯 가지 입력 + 흐름·거버넌스 확인")
    sr.font.size = Pt(12)
    sr.font.bold = True
    sr.font.color.rgb = NAVY

    guide = doc.add_paragraph()
    gr = guide.add_run(
        "V3.1 입력(Q1–Q5): 조직 유형 / 결과물 유형 / 병목 지점 / 현재 팀 규모 / AI 사용 상태. "
        "본 질문지는 견적·일정 확정 전 진단 자료로만 사용됩니다. "
        "실제 개인정보나 내부자료의 원문을 기입하지 마시고, 분류·포함 여부만 표시해 주세요."
    )
    gr.font.size = Pt(9)
    gr.font.color.rgb = GRAY
    gr.font.name = "Malgun Gothic"

    page_header(doc, "Q1–Q5. V3.1 다섯 가지 입력 (필수)", 1)
    choice(doc, "Q1. 조직 유형", ["지역 문화기관", "지역 교육기관", "지역 협회·단체",
                                  "지역 미디어·콘텐츠 기관", "기업 홍보·콘텐츠팀", "기타"])
    short_answer(doc, "Q1-기타", "기타인 경우 자유 기재 (조직명 약칭만, 실명·기밀 금지)")
    choice(doc, "Q2. 결과물 유형", ["홍보물", "교육자료", "영상·이미지", "캠페인 콘텐츠", "보도자료", "기타"])
    short_answer(doc, "Q2-부", "주 결과물 1개 + 부 결과물 목록 (원문 첨부 금지 — 유형만)")
    choice(doc, "Q3. 병목 지점 (1–2개 선택 + 짧은 이유)", ["기획", "초안", "제작", "검토", "승인", "배포"])
    short_answer(doc, "Q3-이유", "선택한 병목의 이유 (인물 비판 대신 프로세스 관점)")

    page_break(doc)

    # ---- Page 2 ----
    page_header(doc, "Q4–Q5. 팀 규모·AI 상태 + Q6–Q9 흐름", 2)
    short_answer(doc, "Q4. 현재 팀 규모", "상시 담당 인원 수 + 역할별 구성 (예: 기획 1 / 제작 2 / 검토 1, 실명 대신 역할·인원만)")
    choice(doc, "Q5. AI 사용 상태", ["미사용", "개인별 탐색", "일부 업무 보조", "승인 도구 운영 중", "기타"])
    short_answer(doc, "Q5-도구", "현재 사용 중인 승인/비승인 도구 간단 목록 (계정 정보 기재 금지)")

    page_header(doc, "Q6–Q9. 흐름·기준선·검토·도구", 2)
    narrative(doc, "Q6. 현재 콘텐츠 제작 흐름", "기획→초안→검토→승인→게시 순으로")
    narrative(doc, "Q7. 업무별 소요시간", "콘텐츠 1건당 단계별 시간 표 (기준 생산시간)")
    narrative(doc, "Q8. 검토·승인 단계", "단계 목록 + 역할(직책만)")
    narrative(doc, "Q9. 현재 사용하는 AI 도구 상세", "Q5 요약과 별개로 도구별 상세 (도구명 + 용도 + 빈도)")

    page_break(doc)

    # ---- Page 3 ----
    page_header(doc, "Q10–Q17. 거버넌스·파일럿 준비", 3)
    choice(doc, "Q10. 개인정보", ["예", "아니오", "일부"])
    choice(doc, "Q11. 저작권 자료", ["예", "아니오", "일부"])
    short_answer(doc, "Q12. 외부 공개 여부", "외부에 공개되는 채널 목록 (웹·SNS·언론 배포 등)")
    narrative(doc, "Q13. 재작업·실패 유형", "빈도 높은 순서 나열 (프로세스 관점)")
    yesno(doc, "Q14. 과거 교육 경험 (선택)", ["예", "아니오"])
    narrative(doc, "Q15. 금지 업무", "AI를 절대 쓰면 안 되는 업무")
    narrative(doc, "Q16. 파일럿 담당자", "1팀 6–10명 + 운영 책임자 1인 (역할·인원·주당 투입 시간, Q4와 별개)")
    short_answer(doc, "Q17. 예산 승인자", "파일럿 예산 최종 승인 직책 (실명 대신 직책만)")

    closing = doc.add_paragraph()
    cr = closing.add_run(
        "\n본 질문지는 견적·일정 확정 전 진단 자료로만 사용됩니다. "
        "가격은 시장 검증 전 자사 가격 가설이며 실제 계약·매출 주장이 아닙니다. "
        "조직별 사용정책·검토체계, 개인정보·저작권·조달 관련 사항은 고객별 확인과 "
        "전문 법률·계약 검토가 필요합니다."
    )
    cr.font.size = Pt(9)
    cr.font.color.rgb = GRAY
    cr.font.name = "Malgun Gothic"

    Path(OUT).parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.core_properties.created = FIXED_DT
        doc.core_properties.modified = FIXED_DT
        doc.core_properties.revision = 1
    except Exception:
        pass
    doc.save(str(OUT))
    from normalize_ooxml import normalize_ooxml
    normalize_ooxml(OUT)
    print(f"saved {OUT}")


if __name__ == "__main__":
    build()
