"""Build Business32_Skill_Discovery_Worksheet.docx (exactly 2 pages).

Page 1: Q1~Q7, page 2: Q8~Q13 with a repeated header on both pages.
Descriptive questions use a question number and answer lines; only the
check-type frequency question (Q12) uses a real checkbox.
"""
import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "Business32_Skill_Discovery_Worksheet.docx",
)

QUESTIONS = [
    ("반복되는 업무는 무엇인가", "업무명과 한 문장 설명"),
    ("누가 실행하는가", "담당자 역할"),
    ("누가 검토·승인하는가", "검토자·승인자 역할"),
    ("입력자료는 무엇인가", "업무 시작에 필요한 자료"),
    ("현재 단계는 무엇인가", "현재 수행 순서"),
    ("가장 오래 걸리는 단계는 무엇인가", "병목 단계"),
    ("실패하면 어떤 문제가 생기는가", "오류·재작업 영향"),
    ("AI에 넣으면 안 되는 정보는 무엇인가", "개인정보·기밀 금지 항목"),
    ("필수 증거는 무엇인가", "결과를 뒷받침할 근거"),
    ("예외는 무엇인가", "평소와 다른 특수 상황"),
    ("최종 승인 기준은 무엇인가", "승인 조건"),
    ("얼마나 자주 반복되는가", "주간·월간 빈도"),
    ("출력 형식은 무엇인가", "최종 산출물 형태"),
]

CHECK_Q = {12}  # check-type question numbers (checkbox prefix)

doc = Document()

for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.1)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

style = doc.styles["Normal"]
style.font.name = "NanumSquareRound"
style.font.size = Pt(10)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "NanumSquareRound")


def para(text, size=10, bold=False, color=None, space_after=4, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    return p


def answer_line():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B7B7B7")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def page_header(page_no):
    para("Business 32 · AI Skill Studio", size=16, bold=True, color=(0x2B, 0x6C, 0xB0), space_after=2)
    para("Skill Discovery Worksheet — 반복업무 1개 선정", size=12, bold=True, space_after=2)
    para("제안검토용 DRAFT · 합성 샘플 기반 · 실제 파일·개인정보는 넣지 마세요",
         size=8, color=(0x6B, 0x72, 0x80), space_after=4)
    pn = doc.add_paragraph()
    pn.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = pn.add_run("페이지 %d/2" % page_no)
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    pn.paragraph_format.space_after = Pt(8)


def question(i, question, hint, lines=4):
    p = doc.add_paragraph()
    prefix = "☐ " if i in CHECK_Q else ""
    run = p.add_run("%sQ%d. %s" % (prefix, i, question))
    run.bold = True
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    h = doc.add_paragraph()
    hr = h.add_run("    (힌트: %s)" % hint)
    hr.font.size = Pt(8)
    hr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    h.paragraph_format.space_after = Pt(1)
    if i in CHECK_Q:
        opts = doc.add_paragraph()
        rr = opts.add_run("☐ 주 1회 미만    ☐ 주 1회    ☐ 주 2~3회    ☐ 매일")
        rr.font.size = Pt(9)
        opts.paragraph_format.space_after = Pt(1)
        extra = doc.add_paragraph()
        er = extra.add_run("기타: ")
        er.font.size = Pt(9)
        extra.paragraph_format.space_after = Pt(1)
        answer_line()
    for _ in range(lines):
        answer_line()


page_header(1)
para("기관·팀 (합성 예시: 가상 기관명 A / 홍보팀)", size=9, color=(0x6B, 0x72, 0x80))
answer_line()

for i in range(1, 8):
    q, hint = QUESTIONS[i - 1]
    question(i, q, hint)

doc.add_page_break()

page_header(2)

for i in range(8, 14):
    q, hint = QUESTIONS[i - 1]
    question(i, q, hint)

para("작성 완료 후 확인", size=11, bold=True, color=(0x2F, 0x85, 0x5A), space_after=4)
for line in [
    "업무명·단계·소요시간·검토 구조만 확인하며, 실제 견적서·내부 문서·개인정보를 수집하지 않습니다.",
    "사람 검토와 승인이 필수입니다. AI가 자동 승인하지 않습니다.",
    "모든 예시는 합성입니다.",
    "선정한 반복업무 1개:",
]:
    para(line, size=9, color=(0x6B, 0x72, 0x80), space_after=2)
answer_line()

doc.save(OUT)
print("saved", OUT)
